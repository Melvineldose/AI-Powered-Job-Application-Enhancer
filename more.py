import streamlit as st
import pandas as pd
import numpy as np
import re
import os
import io
import base64
import time
import uuid
import json
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
from dotenv import load_dotenv
import openai
import PyPDF2
import docx
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import plotly.express as px
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document
import hashlib
import pickle
from matplotlib_venn import venn2

# Load environment variables - this loads the API key from .env file
load_dotenv()

# Initialize OpenAI client with API key from .env
openai.api_key = os.getenv("OPENAI_API_KEY")

# Create cache directory if it doesn't exist
if not os.path.exists(".cache"):
    os.makedirs(".cache")

# Set page configuration
st.set_page_config(
    page_title="Job Description & Resume Matcher",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state for all components
def init_session_state():
    """Initialize all session state variables if they don't exist"""
    if 'optimized_resume' not in st.session_state:
        st.session_state.optimized_resume = None
    if 'resume_text' not in st.session_state:
        st.session_state.resume_text = ""
    if 'job_desc_text' not in st.session_state:
        st.session_state.job_desc_text = ""
    if 'match_result' not in st.session_state:
        st.session_state.match_result = None
    if 'recommendations' not in st.session_state:
        st.session_state.recommendations = None
    if 'analysis_done' not in st.session_state:
        st.session_state.analysis_done = False
    if 'resume_history' not in st.session_state:
        st.session_state.resume_history = []
    if 'theme' not in st.session_state:
        st.session_state.theme = "light"
    if 'cover_letter' not in st.session_state:
        st.session_state.cover_letter = None
    if 'interview_questions' not in st.session_state:
        st.session_state.interview_questions = None
    if 'api_calls_count' not in st.session_state:
        st.session_state.api_calls_count = 0
    if 'last_api_call_time' not in st.session_state:
        st.session_state.last_api_call_time = 0
    if 'custom_weights' not in st.session_state:
        st.session_state.custom_weights = {"skills": 60, "content": 40}
    if 'keyword_analysis' not in st.session_state:
        st.session_state.keyword_analysis = None

# Call init function
init_session_state()

# Apply theme
def apply_theme():
    """Apply selected theme using custom CSS"""
    if st.session_state.theme == "dark":
        st.markdown("""
        <style>
        .main {background-color: #0E1117; color: #FAFAFA;}
        .sidebar .sidebar-content {background-color: #262730; color: #FAFAFA;}
        .stButton>button {background-color: #4CAF50; color: white;}
        .stTextInput>div>div>input {background-color: #262730; color: #FAFAFA;}
        .stTextArea>div>div>textarea {background-color: #262730; color: #FAFAFA;}
        .css-145kmo2 {color: #FAFAFA !important;}
        </style>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <style>
        .stButton>button {background-color: #4CAF50; color: white;}
        </style>
        """, unsafe_allow_html=True)

# Cache management
def get_cache_key(data_str):
    """Generate a unique cache key for a given string input"""
    return hashlib.md5(data_str.encode()).hexdigest()

def save_to_cache(key, data):
    """Save data to cache file"""
    try:
        with open(f".cache/{key}.pkl", "wb") as f:
            pickle.dump(data, f)
        return True
    except Exception as e:
        print(f"Cache save error: {e}")
        return False

def load_from_cache(key):
    """Load data from cache file if it exists"""
    try:
        if os.path.exists(f".cache/{key}.pkl"):
            with open(f".cache/{key}.pkl", "rb") as f:
                return pickle.load(f)
    except Exception as e:
        print(f"Cache load error: {e}")
    return None

# Rate limiting for API calls
def check_rate_limit():
    """Check if we're within rate limits for API calls"""
    # Simple rate limiting: max 50 calls per hour, with at least 2 seconds between calls
    current_time = time.time()
    
    # Check if enough time has passed since last call
    if current_time - st.session_state.last_api_call_time < 2:
        time.sleep(2)  # Wait until we're allowed to make another call
    
    # Check hourly limit
    if st.session_state.api_calls_count >= 50:
        # Reset counter if an hour has passed
        if current_time - st.session_state.last_api_call_time > 3600:
            st.session_state.api_calls_count = 0
        else:
            return False
    
    # Update rate limit tracking
    st.session_state.last_api_call_time = time.time()
    st.session_state.api_calls_count += 1
    return True

# Helper Functions
def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF file"""
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_file):
    """Extract text from uploaded DOCX file"""
    doc = docx.Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text_from_file(file):
    """Extract text from uploaded file based on extension"""
    if file is None:
        return ""
        
    file_extension = file.name.split('.')[-1].lower()
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(file)
    elif file_extension in ['docx', 'doc']:
        return extract_text_from_docx(file)
    elif file_extension in ['txt', 'md']:
        return file.getvalue().decode('utf-8')
    else:
        st.error(f"Unsupported file format: {file_extension}")
        return ""

def clean_text(text):
    """Basic text cleaning without relying on NLTK"""
    # Convert to lowercase
    text = text.lower()
    
    # Remove special characters, numbers, and extra whitespace
    text = re.sub(r'[^a-zA-Z\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

def extract_skills_and_analyze(resume_text, job_text, model="gpt-3.5-turbo", weights=None):
    """Use OpenAI to extract skills and analyze match with option for model selection and weights"""
    # Use default weights if none provided
    if weights is None:
        weights = {"skills": 60, "content": 40}
    
    # Generate cache key based on inputs and model
    cache_key = get_cache_key(f"{resume_text[:1000]}{job_text[:1000]}{model}{weights}")
    cached_result = load_from_cache(cache_key)
    
    if cached_result:
        return cached_result
    
    if not check_rate_limit():
        st.error("API rate limit reached. Please wait before making more requests.")
        return None
    
    try:
        prompt = f"""
        # Resume and Job Description Analysis Task

        ## Resume:
        {resume_text[:3000]}...

        ## Job Description:
        {job_text[:3000]}...

        ## Custom Weights:
        - Skill Match Weight: {weights['skills']}%
        - Content Match Weight: {weights['content']}%

        Please analyze these documents and provide the following in JSON format:
        1. All skills mentioned in the job description (technical, soft skills, certifications)
        2. All skills mentioned in the resume
        3. Common skills found in both
        4. Skills in the job description not found in the resume
        5. A numerical score for skill match (percentage of job skills found in resume)
        6. A numerical score for overall content similarity (semantic match)
        7. A weighted total score using the custom weights provided
        8. Detect the industry and role from the job description

        Format your response as valid JSON with the following structure:
        {{
            "job_skills": ["skill1", "skill2", ...],
            "resume_skills": ["skill1", "skill2", ...],
            "common_skills": ["skill1", "skill2", ...],
            "missing_skills": ["skill1", "skill2", ...],
            "skill_match": percentage_number,
            "content_match": percentage_number,
            "overall_score": percentage_number,
            "industry": "detected_industry",
            "role": "detected_role"
        }}

        Return ONLY the JSON, no other text.
        """

        # Show progress bar
        progress_text = "Running AI analysis..."
        progress_bar = st.progress(0)
        progress_bar.progress(10, text=progress_text)
        
        # Call OpenAI API for analysis
        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert ATS system that analyzes resumes and job descriptions. You extract skills and calculate match scores accurately."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.3
        )
        
        progress_bar.progress(60, text=progress_text)
        
        result = response.choices[0].message.content.strip()
        
        # Try to parse JSON from response
        try:
            # First, try direct JSON parsing
            match_data = json.loads(result)
            
            # Calculate keyword density separately
            keyword_analysis = analyze_keyword_density(resume_text, job_text, match_data["job_skills"])
            match_data["keyword_analysis"] = keyword_analysis
            
            # Cache the result
            save_to_cache(cache_key, match_data)
            
            progress_bar.progress(100, text="Analysis complete!")
            time.sleep(0.5)
            progress_bar.empty()
            
            return match_data
        except json.JSONDecodeError:
            # If direct parsing fails, try to extract JSON from text
            try:
                # Look for JSON content between triple backticks if present
                json_match = re.search(r'```json\s*(.*?)\s*```', result, re.DOTALL)
                if json_match:
                    match_data = json.loads(json_match.group(1))
                    
                    # Calculate keyword density separately
                    keyword_analysis = analyze_keyword_density(resume_text, job_text, match_data["job_skills"])
                    match_data["keyword_analysis"] = keyword_analysis
                    
                    # Cache the result
                    save_to_cache(cache_key, match_data)
                    
                    progress_bar.progress(100, text="Analysis complete!")
                    time.sleep(0.5)
                    progress_bar.empty()
                    
                    return match_data
                
                # Try to find JSON content at the beginning of any line
                json_match = re.search(r'^\s*({.*})\s*$', result, re.MULTILINE | re.DOTALL)
                if json_match:
                    match_data = json.loads(json_match.group(1))
                    
                    # Calculate keyword density separately
                    keyword_analysis = analyze_keyword_density(resume_text, job_text, match_data["job_skills"])
                    match_data["keyword_analysis"] = keyword_analysis
                    
                    # Cache the result
                    save_to_cache(cache_key, match_data)
                    
                    progress_bar.progress(100, text="Analysis complete!")
                    time.sleep(0.5)
                    progress_bar.empty()
                    
                    return match_data
                    
                # Log the response for debugging
                progress_bar.empty()
                st.error("Could not extract valid JSON from the OpenAI response")
                st.code(result)
                return None
            except Exception as inner_e:
                progress_bar.empty()
                st.error(f"Error parsing JSON from response: {str(inner_e)}")
                st.code(result)
                return None
    except Exception as e:
        st.error(f"Error calling OpenAI API: {str(e)}")
        return None

def analyze_keyword_density(resume_text, job_text, keywords):
    """Analyze keyword frequency and density in both resume and job description"""
    result = {}
    
    # Clean texts
    resume_clean = clean_text(resume_text.lower())
    job_clean = clean_text(job_text.lower())
    
    # Count occurrences of each keyword
    for keyword in keywords:
        keyword_clean = clean_text(keyword.lower())
        
        # Skip keywords that are just one character
        if len(keyword_clean) < 2:
            continue
            
        # Count exact matches
        resume_count = len(re.findall(r'\b{}\b'.format(re.escape(keyword_clean)), resume_clean))
        job_count = len(re.findall(r'\b{}\b'.format(re.escape(keyword_clean)), job_clean))
        
        # Calculate density (occurrences per 1000 words)
        resume_words = len(resume_clean.split())
        job_words = len(job_clean.split())
        
        resume_density = (resume_count / resume_words) * 1000 if resume_words > 0 else 0
        job_density = (job_count / job_words) * 1000 if job_words > 0 else 0
        
        # Store results
        result[keyword] = {
            "resume_count": resume_count,
            "job_count": job_count,
            "resume_density": round(resume_density, 2),
            "job_density": round(job_density, 2),
            "density_ratio": round(resume_density / job_density, 2) if job_density > 0 else 0
        }
    
    return result

def calculate_text_similarity(text1, text2):
    """Calculate basic text similarity using TF-IDF and cosine similarity"""
    # Clean texts
    clean_text1 = clean_text(text1)
    clean_text2 = clean_text(text2)
    
    # Calculate text similarity using TF-IDF and cosine similarity
    vectorizer = TfidfVectorizer()
    try:
        tfidf_matrix = vectorizer.fit_transform([clean_text1, clean_text2])
        text_similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0] * 100
        return min(round(text_similarity, 1), 100)
    except:
        return 0  # Return 0 if vectorization fails

def get_gpt_recommendations(resume_text, job_text, match_result, model="gpt-3.5-turbo"):
    """Get recommendations from OpenAI with option for model selection"""
    # Generate cache key
    cache_key = get_cache_key(f"recommendations_{resume_text[:1000]}{job_text[:1000]}{model}")
    cached_result = load_from_cache(cache_key)
    
    if cached_result:
        return cached_result
    
    if not check_rate_limit():
        st.error("API rate limit reached. Please wait before making more requests.")
        return "API rate limit reached. Please try again later."
    
    try:
        # Show progress indicator
        progress_text = "Generating recommendations..."
        progress_bar = st.progress(0)
        progress_bar.progress(20, text=progress_text)
        
        # Create prompt for GPT
        prompt = f"""
        # Resume and Job Description Analysis

        ## Resume:
        {resume_text[:1500]}...

        ## Job Description:
        {job_text[:1500]}...

        ## Match Analysis:
        - Overall Match Score: {match_result.get('overall_score', 0)}%
        - Skill Match: {match_result.get('skill_match', 0)}%
        - Content Match: {match_result.get('content_match', 0)}%
        - Common Skills: {', '.join(match_result.get('common_skills', []))}
        - Missing Skills: {', '.join(match_result.get('missing_skills', []))}
        - Industry: {match_result.get('industry', 'Unknown')}
        - Role: {match_result.get('role', 'Unknown')}

        Based on the above information, please provide:
        1. A summary of the strengths and weaknesses of the resume relative to the job description
        2. 3-5 specific recommendations to improve the resume to better match the job description
        3. 3 bullet points that could be added to the resume to better align with the job requirements
        4. Industry-specific advice for this role
        5. Recommended certifications that could enhance the applicant's profile
        
        Keep your response concise and actionable. Format it with markdown.
        """

        # Call OpenAI API
        progress_bar.progress(40, text=progress_text)
        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert ATS resume optimization assistant helping job seekers tailor their resumes to specific job descriptions."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.5
        )
        
        progress_bar.progress(90, text=progress_text)
        recommendations = response.choices[0].message.content
        
        # Cache the result
        save_to_cache(cache_key, recommendations)
        
        progress_bar.progress(100, text="Recommendations ready!")
        time.sleep(0.5)
        progress_bar.empty()
        
        return recommendations
    except Exception as e:
        st.error(f"Error getting recommendations: {str(e)}")
        return "Error generating recommendations. Please check your OpenAI API key and try again."

def get_interview_questions(resume_text, job_text, match_result, model="gpt-3.5-turbo"):
    """Generate potential interview questions based on the job-resume gap analysis"""
    # Generate cache key
    cache_key = get_cache_key(f"interview_{resume_text[:1000]}{job_text[:1000]}{model}")
    cached_result = load_from_cache(cache_key)
    
    if cached_result:
        return cached_result
    
    if not check_rate_limit():
        st.error("API rate limit reached. Please wait before making more requests.")
        return "API rate limit reached. Please try again later."
    
    try:
        # Create prompt for GPT
        prompt = f"""
        # Interview Question Generation

        ## Resume Summary:
        {resume_text[:1000]}...

        ## Job Description:
        {job_text[:1000]}...

        ## Match Analysis:
        - Missing Skills: {', '.join(match_result.get('missing_skills', []))}
        - Common Skills: {', '.join(match_result.get('common_skills', []))}
        - Industry: {match_result.get('industry', 'Unknown')}
        - Role: {match_result.get('role', 'Unknown')}

        Generate 10 potential interview questions the candidate might face for this position, divided into these categories:
        1. Technical/skill-based questions (focusing on both the skills they have and skills they're missing)
        2. Behavioral questions specific to this role
        3. Questions about the gap between their experience and the job requirements
        4. Industry-specific questions

        For each question, also provide a brief note on what the interviewer is looking for and a sample strong answer the candidate could give.
        Format your response with markdown, using headers to separate the question categories.
        """

        with st.spinner("Generating interview questions..."):
            # Call OpenAI API
            response = openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are an expert career coach who specializes in interview preparation."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.7
            )
            
            questions = response.choices[0].message.content
            
            # Cache the result
            save_to_cache(cache_key, questions)
            
            return questions
    except Exception as e:
        st.error(f"Error generating interview questions: {str(e)}")
        return "Error generating interview questions. Please check your OpenAI API key and try again."

def get_gpt_optimized_resume(resume_text, job_text, match_result, model="gpt-3.5-turbo"):
    """Get an optimized version of the resume from OpenAI with option for model selection"""
    # Generate cache key
    cache_key = get_cache_key(f"optimized_{resume_text[:1000]}{job_text[:1000]}{model}")
    cached_result = load_from_cache(cache_key)
    
    if cached_result:
        return cached_result
    
    if not check_rate_limit():
        st.error("API rate limit reached. Please wait before making more requests.")
        return "API rate limit reached. Please try again later."
    
    try:
        # Show progress indicator
        progress_text = "Optimizing your resume..."
        progress_bar = st.progress(0)
        progress_bar.progress(20, text=progress_text)
        
        # Create prompt for GPT
        prompt = f"""
        # Resume Optimization Task

        ## Original Resume:
        {resume_text[:2500]}...

        ## Job Description:
        {job_text[:1500]}...

        ## Match Analysis:
        - Missing Skills: {', '.join(match_result.get('missing_skills', []))}
        - Industry: {match_result.get('industry', 'Unknown')}
        - Role: {match_result.get('role', 'Unknown')}
        
        Please rewrite the resume to better match the job description while maintaining honesty and the original structure. Specifically:
        
        1. Keep the same general format and sections
        2. Emphasize skills and experiences that align with the job description
        3. Add relevant keywords from the job description where appropriate
        4. Incorporate missing skills where the candidate might reasonably have them
        5. Quantify achievements where possible
        6. Keep the length similar to the original
        7. Format with markdown to preserve structure
        8. Highlight the changes you've made by putting them in **bold**
        
        Return only the optimized resume text, formatted professionally with markdown.
        """

        # Call OpenAI API
        progress_bar.progress(40, text=progress_text)
        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert resume writer who specializes in optimizing resumes for ATS systems while maintaining honesty and professionalism."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.4
        )
        
        progress_bar.progress(90, text=progress_text)
        optimized_resume = response.choices[0].message.content
        
        # Save to history
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        history_item = {
            "timestamp": timestamp,
            "job_title": match_result.get('role', 'Unknown Job'),
            "match_score": match_result.get('overall_score', 0),
            "resume_text": optimized_resume
        }
        st.session_state.resume_history.append(history_item)
        
        # Cache the result
        save_to_cache(cache_key, optimized_resume)
        
        progress_bar.progress(100, text="Resume optimization complete!")
        time.sleep(0.5)
        progress_bar.empty()
        
        return optimized_resume
    except Exception as e:
        st.error(f"Error getting optimized resume: {str(e)}")
        return "Error generating optimized resume. Please check your OpenAI API key and try again."

def generate_cover_letter(resume_text, job_text, match_result, model="gpt-3.5-turbo"):
    """Generate a tailored cover letter based on the resume and job description"""
    # Generate cache key
    cache_key = get_cache_key(f"cover_letter_{resume_text[:1000]}{job_text[:1000]}{model}")
    cached_result = load_from_cache(cache_key)
    
    if cached_result:
        return cached_result
    
    if not check_rate_limit():
        st.error("API rate limit reached. Please wait before making more requests.")
        return "API rate limit reached. Please try again later."
    
    try:
        # Show progress indicator
        progress_text = "Generating cover letter..."
        progress_bar = st.progress(0)
        progress_bar.progress(20, text=progress_text)
        
        # Create prompt for GPT
        prompt = f"""
        # Cover Letter Generation Task

        ## Resume:
        {resume_text[:1500]}...

        ## Job Description:
        {job_text[:1500]}...

        ## Match Analysis:
        - Common Skills: {', '.join(match_result.get('common_skills', []))}
        - Industry: {match_result.get('industry', 'Unknown')}
        - Role: {match_result.get('role', 'Unknown')}
        
        Generate a professional cover letter that:
        1. Follows standard cover letter format with date and proper salutation
        2. Highlights the applicant's relevant skills and experiences from their resume
        3. Addresses how the applicant's background aligns with key requirements in the job description
        4. Demonstrates knowledge of the industry 
        5. Has a confident and enthusiastic tone
        6. Ends with a strong closing paragraph requesting an interview
        7. Is approximately 350-400 words in length
        
        Format the letter in markdown with proper spacing and paragraph structure.
        """

        # Call OpenAI API
        progress_bar.progress(50, text=progress_text)
        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a professional cover letter writer with expertise in creating compelling cover letters tailored to specific job descriptions."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.6
        )
        
        progress_bar.progress(90, text=progress_text)
        cover_letter = response.choices[0].message.content
        
        # Cache the result
        save_to_cache(cache_key, cover_letter)
        
        progress_bar.progress(100, text="Cover letter complete!")
        time.sleep(0.5)
        progress_bar.empty()
        
        return cover_letter
    except Exception as e:
        st.error(f"Error generating cover letter: {str(e)}")
        return "Error generating cover letter. Please check your OpenAI API key and try again."

# Function to generate cover letter when button is clicked
def handle_cover_letter_button():
    with st.spinner("Generating cover letter..."):
        model_choice = st.session_state.get('model_choice', "gpt-3.5-turbo")
        st.session_state.cover_letter = generate_cover_letter(
            st.session_state.resume_text,
            st.session_state.job_desc_text,
            st.session_state.match_result,
            model=model_choice
        )

# Function to generate interview questions when button is clicked
def generate_interview_questions():
    with st.spinner("Generating interview questions..."):
        model_choice = st.session_state.get('model_choice', "gpt-3.5-turbo")
        st.session_state.interview_questions = get_interview_questions(
            st.session_state.resume_text,
            st.session_state.job_desc_text,
            st.session_state.match_result,
            model=model_choice
        )

# Toggle theme
def toggle_theme():
    if st.session_state.theme == "light":
        st.session_state.theme = "dark"
    else:
        st.session_state.theme = "light"
    apply_theme()
def display_match_score(score, is_mobile=False):
    """
    Display a visual representation of the match score with a gauge-like visualization.
    
    Args:
        score (float): The match score percentage (0-100)
        is_mobile (bool): Whether the display is on mobile for responsive design
    """
    # Ensure score is within bounds
    score = max(0, min(100, score))
    
    # Determine color based on score
    if score >= 80:
        color = "green"
        emoji = "🌟"
        assessment = "Excellent Match"
    elif score >= 60:
        color = "orange"
        emoji = "👍"
        assessment = "Good Match"
    elif score >= 40:
        color = "darkorange" 
        emoji = "⚠️"
        assessment = "Fair Match"
    else:
        color = "red"
        emoji = "⚠️"
        assessment = "Poor Match"
    
    # Create gauge visualization
    gauge_html = f"""
    <div style="text-align: center;">
        <div style="margin: 0 auto; width: {'90%' if is_mobile else '60%'};">
            <h1 style="font-size: 4rem; margin-bottom: 0;">{score}%</h1>
            <h2 style="margin-top: 0;">{emoji} {assessment}</h2>
            <div style="background-color: #e0e0e0; height: 30px; border-radius: 15px; margin: 20px 0;">
                <div style="background-color: {color}; width: {score}%; height: 100%; border-radius: 15px; text-align: right;">
                    <span style="color: white; padding-right: 10px; line-height: 30px; font-weight: bold;"></span>
                </div>
            </div>
        </div>
    </div>
    """
    
    # Display score with gauge visualization using HTML
    st.markdown(gauge_html, unsafe_allow_html=True)
    
    # Add interpretation text
    if score >= 80:
        st.success("Your resume is highly aligned with this job description. You're likely to pass ATS screening!")
    elif score >= 60:
        st.info("Your resume shows good alignment with this position. With some targeted improvements, you can increase your chances.")
    elif score >= 40:
        st.warning("Your resume has moderate alignment with this job. Consider significant revisions to improve match rate.")
    else:
        st.error("Your resume has low alignment with this position. Major revisions recommended before applying.")
def highlight_matched_sections(text, keywords):
    """
    Highlight sections of text containing keywords.
    
    Args:
        text (str): Text to highlight keywords in
        keywords (list): List of keywords to highlight
        
    Returns:
        str: HTML formatted text with highlights
    """
    # Clean keywords for better matching
    cleaned_keywords = [clean_text(keyword.lower()) for keyword in keywords]
    
    # Split text into paragraphs
    paragraphs = text.split('\n')
    highlighted_paragraphs = []
    
    for paragraph in paragraphs:
        paragraph_lower = paragraph.lower()
        highlight_paragraph = False
        
        # Check if paragraph contains any keywords
        for keyword in cleaned_keywords:
            if keyword and len(keyword) > 2 and keyword in paragraph_lower:
                highlight_paragraph = True
                break
        
        # Apply highlight style
        if highlight_paragraph and paragraph.strip():
            highlighted_paragraphs.append(f"<div style='background-color: #e6ffe6; padding: 5px; border-left: 3px solid #4CAF50;'>{paragraph}</div>")
        elif paragraph.strip():
            highlighted_paragraphs.append(f"<div>{paragraph}</div>")
        else:
            highlighted_paragraphs.append("<br>")
    
    return "".join(highlighted_paragraphs)

def create_download_link(text, filename, link_text, file_type="txt"):
    """
    Create a download link for different file types.
    
    Args:
        text (str): Text content to download
        filename (str): Name of the file without extension
        link_text (str): Text to display for the download link
        file_type (str): Type of file to generate (txt, md, pdf, docx)
        
    Returns:
        str: HTML link element for downloading content
    """
    if file_type == "pdf":
        # Create a PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Remove markdown formatting and Unicode characters
        clean_content = re.sub(r'(\*\*|__)(.*?)\1', r'\2', text)  # Bold
        clean_content = re.sub(r'(\*|_)(.*?)\1', r'\2', clean_content)  # Italic
        clean_content = re.sub(r'#+\s+', '', clean_content)  # Headers
        
        # Replace problematic Unicode characters with ASCII equivalents
        clean_content = clean_content.replace('\u2013', '-')  # en dash
        clean_content = clean_content.replace('\u2014', '--')  # em dash
        clean_content = clean_content.replace('\u2018', "'")  # single quote start
        clean_content = clean_content.replace('\u2019', "'")  # single quote end
        clean_content = clean_content.replace('\u201c', '"')  # double quote start
        clean_content = clean_content.replace('\u201d', '"')  # double quote end
        clean_content = clean_content.replace('\u2022', '*')  # bullet
        
        # Split by newlines and write to PDF
        for line in clean_content.split('\n'):
            # Try to encode text as Latin-1, replacing characters that can't be encoded
            try:
                line = line.encode('latin-1', errors='replace').decode('latin-1')
                pdf.multi_cell(0, 10, txt=line, align='L')
            except Exception:
                # If it still fails, use a simplified version
                pdf.multi_cell(0, 10, txt="[Text contains unsupported characters]", align='L')
            
        # Save to byte stream - with proper method for BytesIO
        pdf_output = io.BytesIO()
        # Use the correct method for FPDF to write to a bytes buffer
        pdf_bytes = pdf.output(dest='S').encode('latin-1')  # Convert string to bytes
        pdf_output.write(pdf_bytes)
        pdf_output.seek(0)  # Reset pointer to beginning of buffer
        b64 = base64.b64encode(pdf_output.getvalue()).decode()
        
        return f'<a href="data:application/pdf;base64,{b64}" download="{filename}.pdf">{link_text}</a>'
    
    elif file_type == "docx":
        # Create a Word document
        doc = Document()
        
        # Remove markdown formatting
        clean_content = re.sub(r'(\*\*|__)(.*?)\1', r'\2', text)  # Bold
        clean_content = re.sub(r'(\*|_)(.*?)\1', r'\2', clean_content)  # Italic
        
        # Process headers
        lines = clean_content.split('\n')
        for line in lines:
            header_match = re.match(r'^(#+)\s+(.*)', line)
            if header_match:
                level = len(header_match.group(1))
                if level == 1:
                    doc.add_heading(header_match.group(2), level=1)
                elif level == 2:
                    doc.add_heading(header_match.group(2), level=2)
                else:
                    doc.add_heading(header_match.group(2), level=min(level, 9))
            else:
                if line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()
        
        # Save to byte stream
        docx_output = io.BytesIO()
        doc.save(docx_output)
        docx_output.seek(0)  # Reset pointer to beginning of buffer
        b64 = base64.b64encode(docx_output.getvalue()).decode()
        
        return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}.docx">{link_text}</a>'
    
    else:  # txt or md
        b64 = base64.b64encode(text.encode()).decode()
        extension = "md" if file_type == "md" else "txt"
        mime = "text/markdown" if file_type == "md" else "text/plain"
        return f'<a href="data:text/{mime};base64,{b64}" download="{filename}.{extension}">{link_text}</a>'

def generate_optimized_resume():
    """Generate optimized resume and store in session state"""
    with st.spinner("Optimizing your resume..."):
        model_choice = st.session_state.get('model_choice', "gpt-3.5-turbo")
        st.session_state.optimized_resume = get_gpt_optimized_resume(
            st.session_state.resume_text,
            st.session_state.job_desc_text,
            st.session_state.match_result,
            model=model_choice
        )   
def visualize_keyword_density(keyword_analysis):
    """
    Create a visualization of keyword density in resume vs job description.
    
    Args:
        keyword_analysis (dict): Dictionary containing keyword analysis data
        
    Returns:
        fig: Plotly figure object
    """
    # Prepare data for visualization
    keywords = []
    resume_densities = []
    job_densities = []
    
    # Get top keywords by job density
    for keyword, data in sorted(keyword_analysis.items(), 
                               key=lambda x: x[1]['job_density'], 
                               reverse=True)[:15]:  # Top 15 keywords
        keywords.append(keyword)
        resume_densities.append(data['resume_density'])
        job_densities.append(data['job_density'])
    
    # Create grouped bar chart
    fig = go.Figure()
    
    fig.add_trace(go.Bar(
        x=keywords,
        y=resume_densities,
        name='In Your Resume',
        marker_color='#5D9BFB'
    ))
    
    fig.add_trace(go.Bar(
        x=keywords,
        y=job_densities,
        name='In Job Description',
        marker_color='#FD9A6A'
    ))
    
    # Update layout
    fig.update_layout(
        title='Keyword Density Comparison',
        xaxis_title='Keywords',
        yaxis_title='Density (occurrences per 1000 words)',
        barmode='group',
        xaxis={'categoryorder':'total descending'},
        height=500
    )
    
    return fig        
# Main Application
def main():
    # Apply current theme
    apply_theme()
    
    # Set up sidebar
    st.sidebar.image("https://img.icons8.com/color/96/000000/resume.png", width=80)
    st.sidebar.title("Resume ATS Optimizer")
    
    # Add theme toggle
    st.sidebar.button("Toggle Dark/Light Mode", on_click=toggle_theme)
    
    # Check if API key exists in .env
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.sidebar.warning("⚠️ No OpenAI API key found in .env file")
        # API Key input as fallback
        api_key = st.sidebar.text_input("Enter OpenAI API Key", type="password")
        if api_key:
            openai.api_key = api_key
            os.environ["OPENAI_API_KEY"] = api_key
    else:
        st.sidebar.success("✅ OpenAI API key loaded from .env file")
    
    # Model selection
    model_options = {
        "gpt-3.5-turbo": "GPT-3.5 Turbo (Fast & Economic)",
        "gpt-4": "GPT-4 (High Quality, More Expensive)"
    }
    selected_model = st.sidebar.selectbox(
        "Select AI Model", 
        options=list(model_options.keys()),
        format_func=lambda x: model_options[x],
        index=0
    )
    st.session_state.model_choice = selected_model
    
    # Custom weights for analysis
    st.sidebar.subheader("Analysis Weights")
    skills_weight = st.sidebar.slider("Skills Match Weight", 0, 100, st.session_state.custom_weights["skills"], 5)
    content_weight = 100 - skills_weight
    st.sidebar.write(f"Content Match Weight: {content_weight}%")
    st.session_state.custom_weights = {"skills": skills_weight, "content": content_weight}
    
    # History management in sidebar
    if len(st.session_state.resume_history) > 0:
        st.sidebar.markdown("---")
        st.sidebar.subheader("Resume History")
        for i, item in enumerate(st.session_state.resume_history):
            if st.sidebar.button(f"{item['job_title']} ({item['match_score']}%) - {item['timestamp']}", key=f"history_{i}"):
                st.session_state.optimized_resume = item['resume_text']
    
    st.sidebar.markdown("---")
    st.sidebar.subheader("How it works")
    st.sidebar.markdown("""
    1. Upload your resume and the job description
    2. Get a detailed match analysis using AI
    3. Receive personalized recommendations
    4. Generate an optimized resume, cover letter and more!
    """)
    
    # Main content
    st.title("📊 Advanced Job Description & Resume Matcher")
    st.markdown("Upload your resume and a job description to see how well they match and get optimization suggestions.")
    
    # Create two columns for file uploads
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Upload Resume")
        resume_file = st.file_uploader("Choose your resume file", type=["pdf", "docx", "txt"])
        
        if resume_file is not None:
            resume_text = extract_text_from_file(resume_file)
            st.session_state.resume_text = resume_text
            st.success(f"Successfully extracted text from {resume_file.name}")
            
            with st.expander("View Extracted Resume Text"):
                st.text_area("Resume Content", resume_text, height=300)
    
    with col2:
        st.subheader("Enter Job Description")
        job_desc_option = st.radio("Choose input method", ["Upload File", "Paste Text"])
        
        if job_desc_option == "Upload File":
            job_desc_file = st.file_uploader("Choose job description file", type=["pdf", "docx", "txt"])
            if job_desc_file is not None:
                job_desc_text = extract_text_from_file(job_desc_file)
                st.session_state.job_desc_text = job_desc_text
                st.success(f"Successfully extracted text from {job_desc_file.name}")
                
                with st.expander("View Extracted Job Description Text"):
                    st.text_area("Job Description Content", job_desc_text, height=300)
        else:
            job_desc_text = st.text_area("Paste job description here", height=300)
            if job_desc_text:
                st.session_state.job_desc_text = job_desc_text
    
    # Batch processing section
    with st.expander("Batch Processing (Multiple Job Descriptions)"):
        st.write("Upload multiple job descriptions to compare match scores.")
        batch_files = st.file_uploader("Upload multiple job description files", type=["pdf", "docx", "txt"], accept_multiple_files=True)
        
        if batch_files and st.button("Run Batch Analysis") and st.session_state.resume_text:
            batch_results = []
            progress_text = "Processing batch job descriptions..."
            batch_progress = st.progress(0)
            
            for i, job_file in enumerate(batch_files):
                batch_progress.progress((i) / len(batch_files), text=f"{progress_text} ({i+1}/{len(batch_files)})")
                job_text = extract_text_from_file(job_file)
                
                # Quick local text similarity check (faster than API call)
                similarity = calculate_text_similarity(st.session_state.resume_text, job_text)
                
                batch_results.append({
                    "filename": job_file.name,
                    "similarity_score": similarity
                })
            
            batch_progress.progress(1.0, text="Batch processing complete!")
            
            # Display batch results
            st.subheader("Batch Analysis Results")
            batch_df = pd.DataFrame(batch_results).sort_values('similarity_score', ascending=False)
            
            # Create a bar chart
            fig = px.bar(batch_df, x='filename', y='similarity_score', 
                        title="Resume Match Scores Across Multiple Job Descriptions",
                        labels={'similarity_score': 'Match Score (%)', 'filename': 'Job Description'},
                        color='similarity_score',
                        color_continuous_scale='RdYlGn')
            
            fig.update_layout(xaxis_tickangle=-45)
            st.plotly_chart(fig)
            
            # Display as table
            st.table(batch_df)
            
            # Option to set the highest matching job description as current
            if st.button("Use Highest Matching Job Description"):
                best_match = batch_df.iloc[0]
                best_file = [f for f in batch_files if f.name == best_match['filename']][0]
                st.session_state.job_desc_text = extract_text_from_file(best_file)
                st.success(f"Set job description to best match: {best_match['filename']}")
                st.session_state.analysis_done = False  # Reset to trigger new analysis
    
    # Process button
    if st.button("Analyze Match", type="primary"):
        # Check if all required inputs are provided
        if not st.session_state.resume_text:
            st.error("Please upload a resume file.")
        elif not st.session_state.job_desc_text:
            st.error("Please provide a job description.")
        elif not openai.api_key:
            st.error("Please enter your OpenAI API key in the sidebar or add it to your .env file.")
        else:
            with st.spinner("Analyzing your resume and job description using AI..."):
                # Calculate match score using OpenAI
                match_result = extract_skills_and_analyze(
                    st.session_state.resume_text, 
                    st.session_state.job_desc_text,
                    model=st.session_state.model_choice,
                    weights=st.session_state.custom_weights
                )
                
                if match_result is None:
                    st.error("Failed to analyze the documents. Please try again.")
                    return
                
                # If content_match is not provided by the API, calculate it using TF-IDF
                if 'content_match' not in match_result or match_result['content_match'] is None:
                    match_result['content_match'] = calculate_text_similarity(st.session_state.resume_text, st.session_state.job_desc_text)
                
                # Store match result in session state
                st.session_state.match_result = match_result
                st.session_state.keyword_analysis = match_result.get("keyword_analysis", {})
                
                # Get recommendations from OpenAI
                recommendations = get_gpt_recommendations(
                    st.session_state.resume_text, 
                    st.session_state.job_desc_text, 
                    match_result,
                    model=st.session_state.model_choice
                )
                st.session_state.recommendations = recommendations
                
                # Set analysis_done to True
                st.session_state.analysis_done = True
                
                # Force a rerun to show the results
                st.rerun()
    
    # Show results if analysis is done
    if st.session_state.analysis_done:
        # Display results in tabs
        tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
            "Match Score", 
            "Detailed Analysis", 
            "Recommendations", 
            "Optimized Resume",
            "Cover Letter",
            "Interview Prep"
        ])
        
        with tab1:
            st.header("Resume Match Score")
            
            # Check if viewing on mobile
            is_mobile = False
            if 'mobile' in st.session_state:
                is_mobile = st.session_state.mobile
                
            display_match_score(st.session_state.match_result.get("overall_score", 0), is_mobile)
            
            # Display sub-scores
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Skill Match", f"{st.session_state.match_result.get('skill_match', 0)}%")
            with col2:
                st.metric("Content Match", f"{st.session_state.match_result.get('content_match', 0)}%")
                
            # Display role and industry
            if 'role' in st.session_state.match_result:
                st.info(f"Detected Role: {st.session_state.match_result['role']}")
            if 'industry' in st.session_state.match_result:
                st.info(f"Detected Industry: {st.session_state.match_result['industry']}")
        
        with tab2:
            st.header("Detailed Analysis")
            
            # Skill visualization
            job_skills = st.session_state.match_result.get("job_skills", [])
            resume_skills = st.session_state.match_result.get("resume_skills", [])
            common_skills = st.session_state.match_result.get("common_skills", [])
            missing_skills = st.session_state.match_result.get("missing_skills", [])
            
            # Create sets for Venn diagram data
            job_skills_set = set(job_skills)
            resume_skills_set = set(resume_skills)
            
            # Display Venn diagram using matplotlib
            fig, ax = plt.subplots(figsize=(10, 6))
            venn2(subsets=(len(resume_skills_set - job_skills_set), 
                          len(job_skills_set - resume_skills_set), 
                          len(job_skills_set.intersection(resume_skills_set))), 
                 set_labels=('Resume Skills', 'Job Skills'), 
                 ax=ax)
            st.pyplot(fig)
            
            # Keyword density analysis
            if st.session_state.keyword_analysis:
                st.subheader("Keyword Density Analysis")
                fig = visualize_keyword_density(st.session_state.keyword_analysis)
                st.plotly_chart(fig)
            
            # Common skills
            st.subheader("Skills Found in Both")
            if common_skills:
                col_count = 3
                common_skill_cols = st.columns(col_count)
                for i, skill in enumerate(common_skills):
                    common_skill_cols[i % col_count].markdown(f"✅ {skill.title()}")
            else:
                st.warning("No common skills found.")
            
            # Missing skills
            st.subheader("Skills in Job Description Not Found in Resume")
            if missing_skills:
                col_count = 3
                missing_cols = st.columns(col_count)
                for i, skill in enumerate(missing_skills):
                    missing_cols[i % col_count].markdown(f"❌ {skill.title()}")
            else:
                st.success("Your resume contains all skills mentioned in the job description!")
            
            # Resume with highlighted keywords
            st.subheader("Resume with Highlighted Keywords")
            highlighted_resume = highlight_matched_sections(st.session_state.resume_text, common_skills)
            st.markdown(highlighted_resume, unsafe_allow_html=True)
        
        with tab3:
            st.header("Recommendations")
            st.markdown(st.session_state.recommendations)
            
            # Add certificate suggestions based on missing skills
            st.subheader("Certification Suggestions")
            cert_button = st.button("Get Certification Suggestions")
            
            if cert_button:
                missing_skills = st.session_state.match_result.get("missing_skills", [])
                industry = st.session_state.match_result.get("industry", "Unknown")
                
                with st.spinner("Researching relevant certifications..."):
                    cert_prompt = f"""
                    Based on these missing skills: {', '.join(missing_skills)} 
                    and industry: {industry}, suggest 3-5 relevant professional certifications 
                    that would enhance the candidate's profile. For each certification, provide:
                    1. Name of certification
                    2. Issuing organization
                    3. Brief description of what it covers
                    4. Approximate cost and time commitment
                    5. Why it's valuable for this specific role
                    
                    Format as a markdown table.
                    """
                    
                    response = openai.ChatCompletion.create(
                        model=st.session_state.model_choice,
                        messages=[
                            {"role": "system", "content": "You are an expert career advisor specializing in professional certifications."},
                            {"role": "user", "content": cert_prompt}
                        ],
                        max_tokens=1000,
                        temperature=0.5
                    )
                    
                    st.markdown(response.choices[0].message.content)
        
        with tab4:
            st.header("Optimized Resume")
            
            # Create a button that will trigger the generate_optimized_resume function
            if st.button("Generate Optimized Resume", key="gen_resume_btn"):
                generate_optimized_resume()
            
            # Display optimized resume if it exists
            if st.session_state.optimized_resume:
                st.markdown(st.session_state.optimized_resume)
                
                # Create download links for different formats
                st.subheader("Download Options")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown(
                        create_download_link(st.session_state.optimized_resume, "optimized_resume", "Download as Markdown", "md"),
                        unsafe_allow_html=True
                    )
                
                with col2:
                    st.markdown(
                        create_download_link(st.session_state.optimized_resume, "optimized_resume", "Download as PDF", "pdf"),
                        unsafe_allow_html=True
                    )
                
                with col3:
                    st.markdown(
                        create_download_link(st.session_state.optimized_resume, "optimized_resume", "Download as DOCX", "docx"),
                        unsafe_allow_html=True
                    )
                
                # Option to save to history
                if st.button("Save this version to history"):
                    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
                    history_item = {
                        "timestamp": timestamp,
                        "job_title": st.session_state.match_result.get('role', 'Unknown Job'),
                        "match_score": st.session_state.match_result.get('overall_score', 0),
                        "resume_text": st.session_state.optimized_resume
                    }
                    
                    # Check if it's already in history
                    if not any(item["timestamp"] == timestamp for item in st.session_state.resume_history):
                        st.session_state.resume_history.append(history_item)
                        st.success("Saved to history!")
        
        with tab5:
            st.header("Cover Letter Generator")
            
            # Button to generate cover letter
            if st.button("Generate Cover Letter", key="gen_cover_letter_btn"):
                handle_cover_letter_button()
            
            # Display cover letter if it exists
            if st.session_state.cover_letter:
                st.markdown(st.session_state.cover_letter)
                
                # Download options
                st.subheader("Download Options")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown(
                        create_download_link(st.session_state.cover_letter, "cover_letter", "Download as Markdown", "md"),
                        unsafe_allow_html=True
                    )
                
                with col2:
                    st.markdown(
                        create_download_link(st.session_state.cover_letter, "cover_letter", "Download as PDF", "pdf"),
                        unsafe_allow_html=True
                    )
                
                with col3:
                    st.markdown(
                        create_download_link(st.session_state.cover_letter, "cover_letter", "Download as DOCX", "docx"),
                        unsafe_allow_html=True
                    )
        
        with tab6:
            st.header("Interview Preparation")
            
            # Button to generate interview questions
            if st.button("Generate Interview Questions", key="gen_interview_btn"):
                generate_interview_questions()
            
            # Display interview questions if they exist
            if st.session_state.interview_questions:
                st.markdown(st.session_state.interview_questions)
                
                # Download option
                st.markdown(
                    create_download_link(st.session_state.interview_questions, "interview_prep", "Download Interview Prep", "md"),
                    unsafe_allow_html=True
                )
            
            # Mock interview simulator
            st.subheader("Mock Interview Simulator")
            if st.button("Start Mock Interview Simulation"):
                with st.spinner("Setting up mock interview..."):
                    # Generate a set of interview questions specific to the job
                    mock_prompt = f"""
                    Based on this job description:
                    {st.session_state.job_desc_text[:1000]}...
                    
                    And these missing skills:
                    {', '.join(st.session_state.match_result.get('missing_skills', []))}
                    
                    Create a short mock interview scenario with 3 challenging but common interview questions for this role.
                    For each question, provide:
                    1. The interviewer's question
                    2. What the interviewer is looking for in the answer
                    3. A sample strong answer that the candidate could give
                    
                    Format using markdown with clear headings for each question.
                    """
                    
                    response = openai.ChatCompletion.create(
                        model=st.session_state.model_choice,
                        messages=[
                            {"role": "system", "content": "You are an experienced hiring manager conducting a job interview."},
                            {"role": "user", "content": mock_prompt}
                        ],
                        max_tokens=1200,
                        temperature=0.7
                    )
                    
                    st.markdown(response.choices[0].message.content)
                    
                    # Add a tip for a real mock interview
                    st.info("💡 Tip: For a more realistic mock interview experience, ask a friend to use these questions to conduct a practice interview with you.")

    # Add a footer with app information
    st.markdown("---")
    st.markdown(
        """
        <div style="text-align: center;">
            <p>Resume ATS Optimizer • Enhanced Edition • Powered by AI</p>
        </div>
        """,
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()